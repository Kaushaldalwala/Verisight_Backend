from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_documentation():
    doc = Document()

    # Title
    title = doc.add_heading('VeriSight Phase 1 - System Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Overview
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "VeriSight Phase 1 is a comprehensive system designed to automate the extraction and validation of identity documents. "
        "The project successfully merges Optical Character Recognition (OCR) extraction capabilities with a robust Validation Engine (Module 2). "
        "It supports six major document types: Aadhaar, Passport, Driving License, National ID, Permit, and Visa."
    )

    # What We Built
    doc.add_heading('2. What We Have Built', level=1)
    
    p2 = doc.add_paragraph()
    p2.add_run("1. Unified Project Structure:\n").bold = True
    p2.add_run("We combined the heavy OCR machine learning models and the backend Validation engine into a single, cohesive codebase (`versight_phase_1`).\n\n")
    
    p2.add_run("2. Module 2 - Validation Engine API:\n").bold = True
    p2.add_run("We created a standalone API (`POST /api/v1/validation/validate`) that securely takes extracted text (JSON data) from the OCR frontend and validates it. It performs fuzzy name matching, date normalization, and format verification.\n\n")

    p2.add_run("3. Supabase Database Integration:\n").bold = True
    p2.add_run("We successfully connected the Validation Engine to a live Supabase PostgreSQL database. The engine acts as a 'Government Data Provider' by querying actual reference tables (like `val_aadhar_details`, `val_passport_details`) to ensure the OCR data matches ground truth records.\n\n")
    
    p2.add_run("4. Automated Initialization:\n").bold = True
    p2.add_run("We developed a setup script (`setup_db.py`) that automatically creates the required validation tables in Supabase and inserts 300 rows of synthetic mock data (50 for each document type) so the system works out-of-the-box.\n\n")

    p2.add_run("5. Execution Scripts:\n").bold = True
    p2.add_run("A master `run.bat` executable was created so the entire system (virtual environment activation, database initialization, and backend server startup) can be launched with a single click.")

    # Architecture & Flow
    doc.add_heading('3. Integration & API Flow', level=1)
    doc.add_paragraph(
        "The architecture is designed to be highly decoupled and flexible for frontend integration:"
    )
    
    flow_steps = [
        "Step 1: The user uploads an image of an ID document via the Frontend UI.",
        "Step 2: The Frontend sends this image to the OCR API (Module 1).",
        "Step 3: The OCR API extracts the text and returns it to the Frontend as a JSON object.",
        "Step 4: The Frontend forwards this JSON payload directly to the VeriSight Validation API (Module 2).",
        "Step 5: The Validation API queries the Supabase Database to check for a matching ground-truth record.",
        "Step 6: The API returns the final validation status (PASS, FAIL, INCOMPLETE) back to the Frontend."
    ]
    for step in flow_steps:
        doc.add_paragraph(step, style='List Number')

    # Example Payload
    doc.add_heading('4. API Integration Example', level=2)
    doc.add_paragraph("Your frontend or external OCR tool can hit the validation endpoint using standard JSON. For example:")
    
    code_block = doc.add_paragraph(
        'POST /api/v1/validation/validate\n'
        '{\n'
        '  "document_type": "AADHAAR",\n'
        '  "ocr_confidence": 95.5,\n'
        '  "fields": {\n'
        '    "aadhaar_number": "123456789000",\n'
        '    "name": "Test User 0",\n'
        '    "date_of_birth": "12/05/1990"\n'
        '  }\n'
        '}'
    )
    
    # React Frontend Example
    doc.add_heading('5. React Frontend Integration (D:\\...\\clone1\\VeriSight\\frontend)', level=2)
    doc.add_paragraph("If you are integrating with the React/Vite frontend in the clone1 directory, you can create an API utility function like this in your src/api/ folder:")
    
    doc.add_paragraph(
        'export const validateOCRDocument = async (ocrData) => {\n'
        '  try {\n'
        '    const response = await fetch("http://localhost:8000/api/v1/validation/validate", {\n'
        '      method: "POST",\n'
        '      headers: { "Content-Type": "application/json" },\n'
        '      body: JSON.stringify({\n'
        '        document_type: ocrData.docType, // e.g., AADHAAR\n'
        '        ocr_confidence: ocrData.confidence,\n'
        '        fields: ocrData.extractedFields // the JSON object from the OCR API\n'
        '      })\n'
        '    });\n'
        '    const result = await response.json();\n'
        '    return result;\n'
        '  } catch (error) {\n'
        '    console.error("Validation failed", error);\n'
        '  }\n'
        '};'
    )
    
    # Save the document
    file_path = os.path.join(os.path.dirname(__file__), 'VeriSight_Phase1_Documentation.docx')
    doc.save(file_path)
    print(f"Documentation generated successfully at: {file_path}")

if __name__ == "__main__":
    create_documentation()
